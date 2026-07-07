"""Document upload → text extraction → knowledge memory ingestion."""

from __future__ import annotations

import io
import mimetypes
import re
import uuid

from db.database import SessionLocal, get_data_dir
from db.models import Document
from db.persistence import save_message
from memory.fact_extractor import _merge_facts, _parse_extraction_response
from providers.claude_provider import ClaudeProvider

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
MAX_FILE_BYTES = 15 * 1024 * 1024  # 15 MB
MAX_CHARS_FOR_EXTRACTION = 20000  # keep the extraction prompt within a reasonable token budget

# Uploads have no conversation_id of their own; recent-message lookups are global
# (see db.persistence.get_recent_messages_from_db), so any fixed id surfaces here.
DOCUMENT_UPLOAD_CONVERSATION_ID = "document-uploads"

DOCUMENT_EXTRACTION_SYSTEM_PROMPT = """You are extracting useful, reusable facts from a document a user uploaded to an AI companion named Caelo, so Caelo can recall this information in future conversations.

Return ONLY a JSON array of fact objects. No prose, no explanation, just JSON.

Each fact object has this shape:
{"category": "<one of: identity, family, preferences, projects, context, knowledge>", "content": "<short factual statement>"}

Rules:
- Use "knowledge" for facts/information from the document itself (definitions, decisions, data points, plans, reference info)
- Use identity/family/preferences/projects/context only if the document clearly states something about the user
- Each fact should be ONE specific, self-contained statement — split compound facts into separate entries
- Skip boilerplate, formatting artifacts, and anything too vague to be useful later
- If you can't find clear facts, return an empty array []
- Do NOT include markdown code fences in your response — just the raw JSON array"""

DOCUMENT_SUMMARY_SYSTEM_PROMPT = """You are Caelo. A user just uploaded a document — summarize it in your own voice: direct, no padding, no throat-clearing.

Write 3-6 sentences covering what the document actually says — the key points, decisions, or facts a person would want without reading it themselves.

Rules:
- Do NOT start with "Here is a summary" or any preamble — go straight into the content
- Plain prose only — no markdown, no bullet points, no headers
- Be concrete: names, numbers, specifics beat vague generalities
- If the document is thin or mostly boilerplate, say so plainly instead of padding"""

DOCUMENT_CREATION_SYSTEM_PROMPT = """You are Caelo, writing a document a user asked you to create. Produce the document itself — the actual finished content, ready to save and read — not a description of it or notes about writing it.

Rules:
- Output ONLY the document content. No preamble like "Here is the document" and no closing commentary.
- Write in plain text / Markdown. Use Markdown headings and lists only where they genuinely help the document's structure.
- Where the document calls for voice or opinion, it's yours — direct, warm, no corporate filler. Where it's purely factual or practical, just write it cleanly.
- Actually fulfill the request; don't hedge or leave placeholders unless the user explicitly asked for a template."""


class UnsupportedFileType(ValueError):
    pass


class EmptyDocument(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileType(
            f"Unsupported file type {ext or '(none)'!r}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == ".docx":
        from docx import Document

        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
    else:  # pragma: no cover - guarded by the check above
        raise UnsupportedFileType(ext)

    text = text.strip()
    if not text:
        raise EmptyDocument("No extractable text found in this document.")
    return text


def _save_upload_to_disk(filename: str, data: bytes) -> tuple[str, str, int]:
    """Write the raw upload to disk under a generated name. Returns (file_path, content_type, size_bytes)."""
    uploads_dir = get_data_dir() / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    file_path = uploads_dir / f"{uuid.uuid4().hex}{ext}"
    file_path.write_bytes(data)
    content_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
    return str(file_path), content_type, len(data)


def ingest_document(filename: str, data: bytes) -> dict:
    """Extract text from an uploaded document and store distilled facts in long-term memory."""
    text = extract_text(filename, data)
    file_path, content_type, size_bytes = _save_upload_to_disk(filename, data)
    excerpt = text[:MAX_CHARS_FOR_EXTRACTION]

    provider = ClaudeProvider()
    response = provider.generate_messages(
        [
            {"role": "system", "content": DOCUMENT_EXTRACTION_SYSTEM_PROMPT},
            {"role": "user", "content": f"Document: {filename}\n\n{excerpt}\n\nReturn the JSON array of facts now."},
        ]
    )
    facts = _parse_extraction_response(response)
    if not facts:
        ends_clean = response.rstrip().endswith(("```", "]"))
        # Length/ends_clean first, since log viewers often clip the line before the tail is visible.
        print(
            f"[document_service] no facts parsed from {filename!r}: "
            f"length={len(response)} ends_clean={ends_clean} "
            f"head={response[:200]!r} tail={response[-200:]!r}"
        )
    facts = [{**f, "content": f"From {filename}: {f['content']}"} for f in facts]

    summary = provider.generate_messages(
        [
            {"role": "system", "content": DOCUMENT_SUMMARY_SYSTEM_PROMPT},
            {"role": "user", "content": f"Document: {filename}\n\n{excerpt}\n\nSummarize it now."},
        ]
    ).strip()

    session = SessionLocal()
    try:
        saved = _merge_facts(session, facts) if facts else 0
        session.add(
            Document(
                filename=filename,
                facts_saved=saved,
                summary=summary,
                file_path=file_path,
                content_type=content_type,
                size_bytes=size_bytes,
            )
        )
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    save_message(
        DOCUMENT_UPLOAD_CONVERSATION_ID,
        "user",
        f"[Uploaded document: {filename} — Caelo extracted {saved} facts from it into memory]",
    )

    return {
        "filename": filename,
        "chars_extracted": len(text),
        "facts_saved": saved,
        "summary": summary,
    }


def list_documents(limit: int = 100) -> list[dict]:
    """Return past uploads, newest first."""
    session = SessionLocal()
    try:
        rows = (
            session.query(Document)
            .order_by(Document.uploaded_at.desc())
            .limit(limit)
            .all()
        )
        return [
            {
                "id": d.id,
                "filename": d.filename,
                "facts_saved": d.facts_saved,
                "summary": d.summary,
                "content_type": d.content_type,
                "uploaded_at": d.uploaded_at.isoformat() if d.uploaded_at else None,
                "has_file": bool(d.file_path),
            }
            for d in rows
        ]
    except Exception as e:
        print(f"[document_service] list_documents failed: {e}")
        return []
    finally:
        session.close()


def _sanitize_display_filename(title: str) -> str:
    """Turn a free-text title into a safe .md display filename."""
    base = re.sub(r"[^\w\- ]+", "", title).strip() or "document"
    if not base.lower().endswith(".md"):
        base = f"{base}.md"
    return base


def _first_sentences(text: str, count: int = 2) -> str:
    """Cheap summary: the first `count` sentences of body text, skipping Markdown headings."""
    body_lines = [ln.strip() for ln in text.splitlines() if ln.strip() and not ln.lstrip().startswith("#")]
    body = " ".join(body_lines)
    if not body:
        return ""
    sentences = re.split(r"(?<=[.!?])\s+", body)
    return " ".join(sentences[:count]).strip()


def create_document(title: str, instructions: str) -> dict:
    """Have Caelo write a document from a prompt and save it as a real file, like an upload."""
    provider = ClaudeProvider()
    content = provider.generate_messages(
        [
            {"role": "system", "content": DOCUMENT_CREATION_SYSTEM_PROMPT},
            {"role": "user", "content": f"Title: {title}\n\nWhat to write:\n{instructions}"},
        ]
    ).strip()

    filename = _sanitize_display_filename(title)
    data = content.encode("utf-8")

    uploads_dir = get_data_dir() / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    disk_path = uploads_dir / f"{uuid.uuid4().hex}.md"
    disk_path.write_bytes(data)

    summary = _first_sentences(content)

    session = SessionLocal()
    try:
        doc = Document(
            filename=filename,
            facts_saved=0,
            summary=summary,
            file_path=str(disk_path),
            content_type="text/markdown",
            size_bytes=len(data),
        )
        session.add(doc)
        session.commit()
        result = {
            "id": doc.id,
            "filename": doc.filename,
            "summary": doc.summary,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            "has_file": True,
        }
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    return result


def save_image(filename: str, data: bytes) -> dict:
    """Store an uploaded image as a file, with no text/fact extraction."""
    file_path, content_type, size_bytes = _save_upload_to_disk(filename, data)

    session = SessionLocal()
    try:
        doc = Document(
            filename=filename,
            facts_saved=0,
            summary=None,
            file_path=file_path,
            content_type=content_type,
            size_bytes=size_bytes,
        )
        session.add(doc)
        session.commit()
        result = {
            "id": doc.id,
            "filename": doc.filename,
            "content_type": doc.content_type,
            "summary": doc.summary,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            "has_file": True,
        }
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()

    return result


def get_document_file(document_id: int) -> dict | None:
    """Return stored-file info for a single document, or None if missing/never had a file on disk."""
    session = SessionLocal()
    try:
        d = session.query(Document).filter_by(id=document_id).first()
        if not d or not d.file_path:
            return None
        return {
            "filename": d.filename,
            "file_path": d.file_path,
            "content_type": d.content_type,
        }
    except Exception as e:
        print(f"[document_service] get_document_file failed: {e}")
        return None
    finally:
        session.close()
